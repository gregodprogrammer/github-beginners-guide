#!/usr/bin/env python3
"""
Convert all markdown curriculum files to PDF and DOCX.
"""
import os
import sys
import subprocess
from pathlib import Path

# Ensure weasyprint and markdown are available
try:
    import markdown
    from weasyprint import HTML, CSS
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing required packages...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "markdown", "weasyprint", "python-docx"])
    import markdown
    from weasyprint import HTML, CSS
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path("/home/greg/learnkimchi")
OUTPUT_DIR = BASE_DIR / "downloads"
OUTPUT_DIR.mkdir(exist_ok=True)

# CSS for better PDF rendering
PDF_CSS = """
@page { size: A4; margin: 2.5cm; }
body {
    font-family: "DejaVu Sans", "Liberation Sans", "Arial", sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
}
h1 { font-size: 22pt; color: #1a1a2e; border-bottom: 3px solid #16213e; padding-bottom: 0.3em; page-break-before: always; }
h1:first-of-type { page-break-before: auto; }
h2 { font-size: 16pt; color: #16213e; margin-top: 1.5em; border-bottom: 1px solid #ddd; padding-bottom: 0.2em; }
h3 { font-size: 13pt; color: #0f3460; margin-top: 1.2em; }
h4 { font-size: 12pt; color: #333; font-style: italic; }
code {
    font-family: "DejaVu Sans Mono", "Liberation Mono", "Courier New", monospace;
    background: #f4f4f4;
    padding: 0.15em 0.4em;
    border-radius: 3px;
    font-size: 10pt;
}
pre {
    background: #f8f8f8;
    border: 1px solid #e1e1e1;
    border-left: 4px solid #16213e;
    padding: 1em;
    overflow-x: auto;
    font-size: 9.5pt;
    line-height: 1.4;
    page-break-inside: avoid;
}
pre code { background: none; padding: 0; border-radius: 0; }
table {
    width: 100%;
    border-collapse: collapse;
    margin: 1em 0;
    font-size: 10pt;
    page-break-inside: avoid;
}
th, td {
    border: 1px solid #ddd;
    padding: 0.5em;
    text-align: left;
}
th { background: #16213e; color: white; font-weight: bold; }
tr:nth-child(even) { background: #f9f9f9; }
blockquote {
    border-left: 4px solid #e94560;
    margin: 1em 0;
    padding: 0.5em 1em;
    background: #fff5f5;
    font-style: italic;
}
hr { border: none; border-top: 2px solid #ddd; margin: 2em 0; }
ul, ol { margin: 0.5em 0; padding-left: 1.5em; }
li { margin: 0.3em 0; }
details { background: #f0f8ff; padding: 1em; border-radius: 5px; margin: 1em 0; border: 1px solid #b0c4de; }
summary { font-weight: bold; cursor: pointer; color: #16213e; }
"""

def md_to_html(md_path):
    """Convert markdown file to HTML string."""
    text = md_path.read_text(encoding='utf-8')
    # Enable table, fenced_code, toc, and other extensions
    html_body = markdown.markdown(
        text,
        extensions=[
            'tables',
            'fenced_code',
            'toc',
            'nl2br',
            'sane_lists',
        ]
    )
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{md_path.stem}</title>
<style>{PDF_CSS}</style>
</head>
<body>
{html_body}
</body>
</html>"""

def convert_to_pdf(md_path, pdf_path):
    """Convert markdown to PDF."""
    html_content = md_to_html(md_path)
    HTML(string=html_content).write_pdf(str(pdf_path))
    print(f"  PDF: {pdf_path.name}")

def convert_to_docx(md_path, docx_path):
    """Basic markdown-to-DOCX conversion."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')
    
    in_code_block = False
    code_lines = []
    
    for line in lines:
        stripped = line.strip()
        
        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Empty line
        if not stripped:
            continue
        
        # Horizontal rule
        if stripped == '---' or stripped.startswith('---'):
            doc.add_paragraph()
            continue
        
        # Headers
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### ') and not stripped.startswith('#### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        # Numbered lists
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == '.':
            p = doc.add_paragraph(stripped[3:], style='List Number')
        # Blockquote
        elif stripped.startswith('>'):
            text_content = stripped[1:].strip() if len(stripped) > 1 else ''
            p = doc.add_paragraph(text_content)
            p.paragraph_format.left_indent = Inches(0.25)
            if p.runs:
                p.runs[0].italic = True
        # Normal paragraph
        else:
            # Very basic inline code handling
            if '`' in stripped:
                p = doc.add_paragraph()
                parts = stripped.split('`')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        run = p.add_run(part)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            else:
                doc.add_paragraph(stripped)
    
    doc.save(str(docx_path))
    print(f"  DOCX: {docx_path.name}")

# File order for combined PDF
combined_files = [
    BASE_DIR / "README.md",
    BASE_DIR / "git-github-devops-bible.md",
    BASE_DIR / "approach-1-linear" / "curriculum.md",
    BASE_DIR / "approach-2-project-driven" / "curriculum.md",
    BASE_DIR / "approach-3-story-mode" / "curriculum.md",
    BASE_DIR / "student-handout-with-solutions.md",
    BASE_DIR / "quizzes-and-assessments.md",
]

print("=" * 50)
print("Converting Git & GitHub DevOps materials")
print("=" * 50)

# Convert individual files
print("\n1. Converting individual files...")
for md_file in combined_files:
    if md_file.exists():
        # Include parent directory name for files named the same (e.g., curriculum.md)
        if md_file.name == 'curriculum.md' and md_file.parent.name not in ('.', ''):
            base_name = f"{md_file.parent.name}"
        else:
            base_name = md_file.stem
        pdf_path = OUTPUT_DIR / f"{base_name}.pdf"
        docx_path = OUTPUT_DIR / f"{base_name}.docx"
        try:
            convert_to_pdf(md_file, pdf_path)
            convert_to_docx(md_file, docx_path)
        except Exception as e:
            print(f"  ERROR converting {md_file.name}: {e}")
    else:
        print(f"  WARNING: {md_file} not found")

# Create combined PDF
print("\n2. Creating combined master PDF...")
try:
    combined_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Git & GitHub for DevOps - Complete Package</title>
<style>{PDF_CSS}</style>
</head>
<body>
"""
    for md_file in combined_files:
        if md_file.exists():
            text = md_file.read_text(encoding='utf-8')
            html_body = markdown.markdown(
                text,
                extensions=['tables', 'fenced_code', 'toc', 'nl2br', 'sane_lists']
            )
            combined_html += f"<div style='page-break-before: always;'></div>"
            combined_html += html_body
            combined_html += "\n<hr>\n"
    
    combined_html += "</body></html>"
    
    combined_pdf = OUTPUT_DIR / "Git-and-GitHub-for-DevOps-Complete-Package.pdf"
    HTML(string=combined_html).write_pdf(str(combined_pdf))
    print(f"  MASTER PDF: {combined_pdf.name}")
except Exception as e:
    print(f"  ERROR creating combined PDF: {e}")

# Create combined DOCX
print("\n3. Creating combined master DOCX...")
try:
    master_docx = OUTPUT_DIR / "Git-and-GitHub-for-DevOps-Complete-Package.docx"
    master = Document()
    style = master.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title page
    title = master.add_heading('Git & GitHub for DevOps', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = master.add_paragraph('Complete Teaching Package')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.italic = True
    master.add_paragraph()
    info = master.add_paragraph('Includes: Curricula · Slides · Exercises · Quizzes · Reference Manual')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    master.add_page_break()
    
    for md_file in combined_files:
        if md_file.exists():
            master.add_heading(md_file.stem.replace('-', ' ').title(), level=1)
            text = md_file.read_text(encoding='utf-8')
            lines = text.split('\n')
            in_code_block = False
            code_lines = []
            
            for line in lines:
                stripped = line.strip()
                if stripped.startswith('```'):
                    if in_code_block:
                        if code_lines:
                            p = master.add_paragraph()
                            run = p.add_run('\n'.join(code_lines))
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                            p.paragraph_format.left_indent = Inches(0.25)
                            p.paragraph_format.space_after = Pt(6)
                        code_lines = []
                        in_code_block = False
                    else:
                        in_code_block = True
                    continue
                if in_code_block:
                    code_lines.append(line)
                    continue
                if not stripped:
                    continue
                if stripped.startswith('# ') and not stripped.startswith('## '):
                    master.add_heading(stripped[2:], level=1)
                elif stripped.startswith('## ') and not stripped.startswith('### '):
                    master.add_heading(stripped[3:], level=2)
                elif stripped.startswith('### ') and not stripped.startswith('#### '):
                    master.add_heading(stripped[4:], level=3)
                elif stripped.startswith('#### '):
                    master.add_heading(stripped[5:], level=4)
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    master.add_paragraph(stripped[2:], style='List Bullet')
                elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == '.':
                    master.add_paragraph(stripped[3:], style='List Number')
                elif stripped.startswith('>'):
                    p = master.add_paragraph(stripped[1:].strip())
                    p.paragraph_format.left_indent = Inches(0.25)
                else:
                    if '`' in stripped:
                        p = master.add_paragraph()
                        parts = stripped.split('`')
                        for i, part in enumerate(parts):
                            if i % 2 == 0:
                                p.add_run(part)
                            else:
                                run = p.add_run(part)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
                    else:
                        master.add_paragraph(stripped)
            master.add_page_break()
    
    master.save(str(master_docx))
    print(f"  MASTER DOCX: {master_docx.name}")
except Exception as e:
    print(f"  ERROR creating combined DOCX: {e}")

print("\n" + "=" * 50)
print("Done! Files in /home/greg/learnkimchi/downloads/")
print("=" * 50)
for f in sorted(OUTPUT_DIR.iterdir()):
    size = f.stat().st_size / 1024
    print(f"  {f.name:<55} {size:>8.1f} KB")
